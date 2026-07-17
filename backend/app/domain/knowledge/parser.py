"""Document parsers: extract plain text from PDF/Word/MD/TXT files."""
from __future__ import annotations

import base64
import functools
from collections.abc import Callable
from pathlib import Path
from typing import TypeVar

import docx  # python-docx
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.domain.exceptions import DocumentParseError
from app.services.multimodal.image_describe import describe_image

T = TypeVar("T")


def _check_exists(path: str) -> Path:
    """Verify file exists or raise DocumentParseError."""
    p = Path(path)
    if not p.exists():
        raise DocumentParseError(doc_id=path, file_path=path, reason="file not found")
    return p


def _handle_parse_errors(
    fn: Callable[[str], T],
) -> Callable[[str], T]:
    """Decorator: wrap non-DocumentParseError exceptions as DocumentParseError.

    Also raises DocumentParseError if the file is missing. The wrapped
    function should accept a file path and return extracted text.
    """

    @functools.wraps(fn)
    def wrapper(path: str) -> T:
        _check_exists(path)
        try:
            return fn(path)
        except DocumentParseError:
            raise
        except Exception as e:  # noqa: BLE001
            raise DocumentParseError(
                doc_id=path, file_path=path, reason=str(e)
            ) from e

    return wrapper


def _read_text_file(path: str) -> str:
    """Read a UTF-8 text file. Falls back to GBK (Chinese Windows)."""
    try:
        return Path(path).read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return Path(path).read_text(encoding="gbk")


@_handle_parse_errors
def parse_txt(path: str) -> str:
    """Parse a plain text file."""
    return _read_text_file(path)


@_handle_parse_errors
def parse_md(path: str) -> str:
    """Parse a Markdown file. Preserves Markdown syntax for LLM consumption."""
    return _read_text_file(path)


@_handle_parse_errors
def parse_pdf(path: str) -> str:
    """Extract text from a PDF file. Joins pages with double newlines.

    One bad page does not fail the whole PDF; the failure is recorded
    inline in the output.
    """
    try:
        reader = PdfReader(path)
    except PdfReadError as e:
        raise DocumentParseError(
            doc_id=path, file_path=path, reason=f"corrupted PDF: {e}"
        ) from e

    texts: list[str] = []
    for page_idx, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
        except Exception as e:  # noqa: BLE001
            page_text = f"[page extraction failed: {e}]"
        image_descs: list[str] = []
        try:
            for img_idx, img in enumerate(getattr(page, "images", []) or []):
                try:
                    img_mime = img.image_type
                    img_mime = (
                        "image/" + img_mime[1:]
                        if img_mime.startswith("/")
                        else (
                            img_mime
                            if img_mime.startswith("image/")
                            else "image/" + img_mime
                        )
                    )
                    b64 = base64.b64encode(img.data).decode("ascii")
                    desc = describe_image(
                        b64, img_mime, f"PDF第{page_idx + 1}页图{img_idx + 1}"
                    )
                    if desc:
                        image_descs.append(desc)
                except Exception:  # noqa: BLE001
                    continue
        except Exception:  # noqa: BLE001
            pass
        merged = page_text + ("\n" + "\n".join(image_descs) if image_descs else "")
        if merged.strip():
            texts.append(merged)
    return "\n\n".join(texts)


@_handle_parse_errors
def parse_docx(path: str) -> str:
    """Extract text from a .docx file.

    Concatenates body paragraphs + flattened table rows with double newlines.
    """
    document = docx.Document(path)
    parts: list[str] = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    parts.extend(_describe_docx_images(document))
    return "\n\n".join(parts)


def _describe_docx_images(document) -> list[str]:
    """从 docx rels 抽出图片,OCR+VLM 描述。失败静默,返回描述列表。"""
    descs: list[str] = []
    try:
        rels = document.part.rels
    except Exception:  # noqa: BLE001
        return descs
    i = 0
    for _rid, rel in rels.items():
        if "image" not in getattr(rel, "reltype", ""):
            continue
        try:
            blob = rel.target_part.blob
            mime = rel.target_part.content_type or "image/png"
            i += 1
            b64 = base64.b64encode(blob).decode("ascii")
            desc = describe_image(b64, mime, f"DOCX图片{i}")
            if desc:
                descs.append(desc)
        except Exception:  # noqa: BLE001
            continue
    return descs
